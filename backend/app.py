import os
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
import re
from io import BytesIO
from captcha.image import ImageCaptcha
import random
import string
from flask import Flask, request, jsonify, send_file, Response, session
from flask_cors import CORS
from decouple import config
import requests
from serpapi import GoogleSearch
import json
from datetime import datetime
import openai
from typing import List, Dict
from scholarly import scholarly
import semanticscholar as sch
from habanero import Crossref
import logging
import httpx
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import os
from flask_session import Session
import secrets

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    current_dir = Path(__file__).resolve().parent
    env_file = current_dir / '.env'
    load_dotenv(env_file)
except:
    # 如果加载 .env 失败，使用默认值
    os.environ.setdefault('FLASK_APP', 'app.py')
    # 设置其他必要的环境变量...

app = Flask(__name__)
app.secret_key = secrets.token_hex(16)

# 配置 Session
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)

# 配置 CORS
CORS(app, resources={
    r"/api/*": {
        "origins": ["http://localhost:3000"],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"],
        "supports_credentials": True
    }
})

# 配置速率限制
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["100 per hour"]
)

# 配置API密钥和代理
GOOGLE_SCHOLAR_API_KEY = config('GOOGLE_SCHOLAR_API_KEY')
SERPAPI_API_KEY = config('SERPAPI_API_KEY', default='')
PROXY_URL = config('PROXY_URL', default='http://api.wlai.vip')
OPENAI_API_KEY = config('OPENAI_API_KEY')
OPENAI_API_BASE = config('OPENAI_API_BASE')

# 配置OpenAI
openai.api_key = OPENAI_API_KEY
openai.api_base = OPENAI_API_BASE

# 可用的模型配置
AVAILABLE_MODELS = {
    'gpt-4o-mini': {
        'name': 'GPT-4o Mini',
        'model_name': 'gpt-4-0125-preview',
        'max_tokens': 128000,
        'temperature': 0.7
    },
    'claude-3-5-haiku-20241022': {
        'name': 'Claude 3.5 Haiku',
        'model_name': 'claude-3-haiku-20240307',
        'max_tokens': 200000,
        'temperature': 0.7
    },
    'gemini-2.0-flash-thinking-exp-1219': {
        'name': 'Gemini 2.0 Flash Thinking',
        'model_name': 'gemini-2.0-flash-thinking-exp-1219',
        'max_tokens': 100000,
        'temperature': 0.7
    }
}

@app.route('/')
def index():
    """根路由，返回API状态信息"""
    return jsonify({
        'status': 'running',
        'version': '1.0',
        'endpoints': {
            'health_check': '/api/health',
            'generate_paper': '/api/generate-paper'
        }
    })

def search_crossref(query: str, max_results: int = 5) -> List[Dict]:
    """使用Crossref API搜索文献"""
    try:
        logger.info(f"Searching Crossref for: {query}")
        url = "https://api.crossref.org/works"
        params = {
            "query": query,
            "rows": max_results,
            "sort": "relevance",
            "select": "DOI,title,author,published-print,container-title",
            "mailto": "your-email@example.com"  # 添加邮箱以获得更好的服务
        }
        
        response = requests.get(url, params=params)
        papers = []
        
        if response.status_code == 200:
            data = response.json()
            for work in data['message']['items']:
                if work.get('title') and work.get('author'):
                    paper_info = {
                        'title': work['title'][0],
                        'authors': [
                            f"{author.get('given', '')} {author.get('family', '')}"
                            for author in work.get('author', [])
                        ],
                        'year': work.get('published-print', {}).get('date-parts', [['']])[0][0],
                        'journal': work.get('container-title', [''])[0] if work.get('container-title') else '',
                        'doi': work.get('DOI', ''),
                        'source': 'Crossref'
                    }
                    papers.append(paper_info)
                    
        logger.info(f"Found {len(papers)} papers from Crossref")
        return papers
    except Exception as e:
        logger.error(f"Crossref search error: {str(e)}")
        return []

def search_semantic_scholar(query: str, max_results: int = 5) -> List[Dict]:
    """使用Semantic Scholar API搜索文献"""
    try:
        logger.info(f"Searching Semantic Scholar for: {query}")
        # 使用 requests 直接调用 API
        url = f"https://api.semanticscholar.org/graph/v1/paper/search"
        params = {
            "query": query,
            "limit": max_results,
            "fields": "title,authors,year,journal,venue,publicationVenue,citationCount,openAccessPdf"
        }
        
        response = requests.get(url, params=params)
        if response.status_code == 200:
            data = response.json()
            papers = []
            
            for paper in data.get('data', []):
                paper_info = {
                    'title': paper.get('title', ''),
                    'authors': [author.get('name', '') for author in paper.get('authors', [])],
                    'year': paper.get('year'),
                    'journal': (paper.get('publicationVenue', {}) or {}).get('name', ''),
                    'doi': paper.get('doi', ''),
                    'source': 'Semantic Scholar'
                }
                papers.append(paper_info)
            
            logger.info(f"Found {len(papers)} papers from Semantic Scholar")
            return papers
        else:
            logger.error(f"Semantic Scholar API error: {response.status_code}")
            return []
            
    except Exception as e:
        logger.error(f"Semantic Scholar search error: {str(e)}")
        return []

def search_google_scholar(query: str, max_results: int = 5) -> List[Dict]:
    """使用Google Scholar API搜索文献"""
    try:
        logger.info(f"Searching Google Scholar for: {query}")
        if not SERPAPI_API_KEY:
            logger.warning("SERPAPI_API_KEY not configured")
            return []
            
        search_params = {
            "engine": "google_scholar",
            "q": query,
            "api_key": SERPAPI_API_KEY,
            "num": max_results,
            "as_ylo": 2000  # 限制从2000年开始的文献
        }
        
        if PROXY_URL:
            search_params["proxy"] = PROXY_URL
        
        search = GoogleSearch(search_params)
        results = search.get_dict()
        papers = []
        
        if 'organic_results' in results:
            for result in results['organic_results']:
                paper_info = {
                    'title': result.get('title', ''),
                    'authors': [author.strip() for author in result.get('authors', '').split(',') if author.strip()],
                    'year': result.get('year', ''),
                    'journal': result.get('publication', ''),
                    'source': 'Google Scholar',
                    'link': result.get('link', '')
                }
                if paper_info['authors'] and paper_info['title']:  # 只添加有作者和标题的结果
                    papers.append(paper_info)
                    
        logger.info(f"Found {len(papers)} papers from Google Scholar")
        return papers
    except Exception as e:
        logger.error(f"Google Scholar search error: {str(e)}")
        return []

def format_chinese_name(name: str) -> str:
    """处理中文姓名格式"""
    # 移除空格
    name = name.replace(' ', '')
    # 检查是否为中文名字
    if re.match(r'^[\u4e00-\u9fff]{2,}$', name):
        # 如果是两个字的名字，第一个字为姓
        if len(name) == 2:
            return f"{name[0]}{name[1]}"
        # 如果是三个字或以上的名字，第一个字为姓
        else:
            return f"{name[0]}{name[1:]}"
    return name

def format_apa7_citation(paper: Dict, for_reference: bool = True) -> Dict:
    """将论文信息转换为APA第7版格式"""
    authors = paper.get('authors', [])
    formatted_authors = []
    
    for author in authors:
        if re.match(r'^[\u4e00-\u9fff]{2,}$', author.replace(' ', '')):
            formatted_name = format_chinese_name(author)
            formatted_authors.append(formatted_name)
        else:
            formatted_authors.append(author)
    
    if len(formatted_authors) > 1:
        author_text = f"{', '.join(formatted_authors[:-1])}, & {formatted_authors[-1]}"
    else:
        author_text = formatted_authors[0] if formatted_authors else "Unknown"
    
    year = paper.get('year', 'n.d.')
    title = paper.get('title', '')
    journal = paper.get('journal', '')
    doi = paper.get('doi', '')
    source = paper.get('source', '')
    citation_id = f"ref-{hash(title) & 0xffffffff}"
    
    if for_reference:
        # 用于参考文献列表的完整格式
        citation = f"{author_text}. ({year}). {title}. "
        if journal:
            citation += f"{journal}. "
        if doi:
            citation += f"[DOI: {doi}](https://doi.org/{doi})"
    else:
        # 用于正文引用的简短格式
        citation = f"[({author_text}, {year})](#{citation_id})"
    
    return {
        'text': citation,
        'doi': doi,
        'url': f"https://doi.org/{doi}" if doi else None,
        'source': source,
        'id': citation_id,
        'short_citation': f"({author_text}, {year})"
    }

def generate_markdown_citation(citation_id: str, authors: str, year: str) -> str:
    """生成带有链接的引用文本"""
    return f"[({authors}, {year})](#{citation_id})"

def normalize_year(year_str):
    """标准化年份格式"""
    if not year_str:
        return 0
    try:
        # 尝试将年份转换为整数
        if isinstance(year_str, int):
            return year_str
        # 处理字符串格式的年份
        year_str = str(year_str).strip()
        # 提取第一个数字序列
        match = re.search(r'\d{4}', year_str)
        if match:
            return int(match.group())
        return 0
    except (ValueError, TypeError):
        return 0

def search_all_sources(query: str, ref_count: int = 10) -> List[Dict]:
    """从所有来源搜索文献"""
    logger.info(f"Searching for {ref_count} papers...")
    
    # 为每个来源分配搜索数量
    source_count = ref_count + 5  # 每个来源多搜索一些，以确保有足够的结果
    
    results = {
        'crossref': search_crossref(query, max_results=source_count),
        'semantic_scholar': search_semantic_scholar(query, max_results=source_count),
        'google_scholar': search_google_scholar(query, max_results=source_count)
    }
    
    # 合并所有来源的结果
    all_papers = []
    for papers in results.values():
        all_papers.extend(papers)
    
    # 去重（基于标题）
    seen_titles = set()
    unique_papers = []
    for paper in all_papers:
        title = paper.get('title', '').lower()
        if title and title not in seen_titles:
            seen_titles.add(title)
            # 确保年份格式一致
            paper['year'] = normalize_year(paper.get('year', '0'))
            unique_papers.append(paper)
    
    # 按年份排序（使用整数比较）
    unique_papers.sort(key=lambda x: x.get('year', 0), reverse=True)
    
    # 确保返回指定数量的文献
    result_papers = unique_papers[:ref_count]
    logger.info(f"Found {len(result_papers)} papers after filtering")
    
    if len(result_papers) < ref_count:
        logger.warning(f"Could not find enough papers. Requested: {ref_count}, Found: {len(result_papers)}")
    
    return result_papers

@app.route('/api/models', methods=['GET'])
def get_models():
    """获取可用的模型列表"""
    return jsonify({
        'success': True,
        'models': [{
            'id': model_id,
            'name': model_info['name']
        } for model_id, model_info in AVAILABLE_MODELS.items()]
    })

def calculate_max_tokens(word_count: int, output_language: str) -> int:
    """计算所需的最大token数"""
    # 中文每个字约占2个token，英文每个单词约占1.5个token
    multiplier = 2 if output_language == 'zh' else 1.5
    # 添加50%的余量确保生成足够的内容
    return int(word_count * multiplier * 1.5)

@app.route('/api/generate-paper', methods=['POST', 'OPTIONS'])
@limiter.limit("10 per minute")
def generate_paper():
    if request.method == 'OPTIONS':
        return '', 204
        
    try:
        data = request.json
        topic = data.get('text', '')
        model_id = data.get('model', 'gpt-4o-mini')
        try:
            word_count = int(data.get('wordCount', 5000))
        except (ValueError, TypeError):
            word_count = 5000
            
        ref_language = data.get('refLanguage', 'en')
        output_language = data.get('outputLanguage', 'en')
        chat_history = data.get('chatHistory', [])
        ref_count = data.get('refCount', 10)  # 获取参考文献数量参数
        
        logger.info(f"Received paper generation request for topic: {topic} using model: {model_id}")
        
        if not topic:
            return jsonify({'success': False, 'error': '请输入论文主题'}), 400
            
        if model_id not in AVAILABLE_MODELS:
            return jsonify({'success': False, 'error': f'不支持的模型: {model_id}'}), 400

        # 添加更详细的日志
        logger.info("Searching for papers...")
        all_papers = search_all_sources(topic, ref_count=ref_count)  # 使用用户指定的数量
        
        total_results = len(all_papers)
        logger.info(f"Found {total_results} papers")
        
        if total_results == 0:
            return jsonify({
                'success': False, 
                'error': '未找到相关文献，请尝试使用不同的关键词'
            }), 400

        # 准备引用
        all_citations = []
        formatted_citations = []
        for paper in all_papers:
            citation_info = format_apa7_citation(paper)
            all_citations.append(citation_info['text'])
            formatted_citations.append(citation_info)

        # 根据语言筛选文献
        if ref_language != 'all':
            all_papers = [p for p in all_papers if is_language_match(p, ref_language)]
        
        # 根据输出语言修改 prompt
        template = data.get('template', '')  # 获取模板
        output_lang_text = "英文" if output_language == "en" else "中文"
        
        prompt = f"""作为一个国际顶尖期刊的学术论文写作助手，请根据主题"{topic}"用{output_lang_text}按{template}要求生成内容。

要求：
1. 引用规范：
    - 正文引用使用APA第7版格式
    - 中文作者姓在前名在后
    - 多作者使用 ", " 连接
    - 参考文献列表的顺序需要按作者姓氏的首字母进行排序

2. 写作要求：
    - 保持学术严谨性
    - 确保逻辑连贯
    - Literature Review 需要分析说明所有搜索到的参考文献
    - 全文的总字数控制在{word_count}词左右
    - 至少引用{ref_count}个参考文献

可用的参考文献：
{chr(10).join(f"    {i+1}. {citation}" for i, citation in enumerate(all_citations))}"""

        logger.info("Generating paper with AI model...")
        def generate():
            try:
                yield json.dumps({'status': 'searching'}) + '\n'
                
                # 搜索文献
                papers = search_all_sources(topic, ref_count=ref_count)
                
                yield json.dumps({'status': 'writing'}) + '\n'
                
                # 生成论文
                client = openai.OpenAI(
                    api_key=OPENAI_API_KEY,
                    base_url=OPENAI_API_BASE,
                    http_client=httpx.Client(
                        timeout=120.0,
                        follow_redirects=True
                    )
                )
                
                response = client.chat.completions.create(
                    model=AVAILABLE_MODELS[model_id]['model_name'],
                    messages=[
                        {"role": "system", "content": "你是一个专业的学术论文写作助手。请确保生成的内容严格符合要求的字数。"},
                        *[{"role": msg["role"], "content": msg["content"]} for msg in chat_history],
                        {"role": "user", "content": prompt}
                    ],
                    temperature=AVAILABLE_MODELS[model_id]['temperature'],
                    max_tokens=calculate_max_tokens(word_count, output_language),
                    stream=True
                )
                
                content = []
                for chunk in response:
                    if chunk.choices[0].delta.content:
                        content.append(chunk.choices[0].delta.content)
                        yield json.dumps({
                            'status': 'generating',
                            'content': ''.join(content)
                        }) + '\n'
                
                # 最终输出
                final_content = ''.join(content)
                yield json.dumps({
                    'status': 'complete',
                    'success': True,
                    'paper': final_content,
                    'references': formatted_citations
                }) + '\n'
                
            except Exception as e:
                logger.error(f"Generation error: {str(e)}", exc_info=True)
                yield json.dumps({
                    'status': 'error',
                    'success': False,
                    'error': str(e)
                }) + '\n'
        
        return Response(generate(), mimetype='text/event-stream')
            
    except Exception as e:
        logger.error(f"General error: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'error': f'服务器错误: {str(e)}'
        }), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'timestamp': datetime.now().isoformat()})

@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    try:
        data = request.json
        content = data.get('content', '')
        references = data.get('references', [])
        filename = data.get('filename', 'paper').strip()[:50]  # 限制文件名长度
        
        doc = Document()
        
        # 设置基本样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # 添加内容
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                doc.add_paragraph(section.strip())
        
        # 添加参考文献
        if references:
            doc.add_heading('References', level=1)
            for ref in references:
                p = doc.add_paragraph()
                p.add_run(ref['text'])
                if ref['url']:
                    p.add_run(f" (DOI: {ref['url']})").italic = True
        
        # 保存到内存
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
        
    except Exception as e:
        logger.error(f"Docx generation error: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500

def is_language_match(paper: Dict, target_language: str) -> bool:
    """判断文献语言是否匹配"""
    title = paper.get('title', '')
    if target_language == 'zh':
        return bool(re.search(r'[\u4e00-\u9fff]', title))
    elif target_language == 'en':
        return not bool(re.search(r'[\u4e00-\u9fff]', title))
    return True

@app.route('/api/captcha', methods=['GET'])
def generate_captcha():
    try:
        image = ImageCaptcha(width=120, height=40)
        code = ''.join(random.choices(string.digits, k=4))
        image_data = image.generate(code)
        
        # 存储验证码到 session
        session['captcha_code'] = code
        
        response = send_file(
            BytesIO(image_data.getvalue()),
            mimetype='image/png'
        )
        return response
    except Exception as e:
        logger.error(f"Captcha generation error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/verify-captcha', methods=['POST'])
def verify_captcha():
    try:
        data = request.json
        user_code = data.get('code', '')
        stored_code = session.get('captcha_code', '')
        
        if user_code and user_code == stored_code:
            session.pop('captcha_code', None)  # 使用后删除
            return jsonify({'success': True})
        return jsonify({'success': False})
    except Exception as e:
        logger.error(f"Captcha verification error: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 